"""PDF and DOCX transcript export.

PDF uses a system Unicode TTF when available (Windows Arial Unicode / Segoe /
Malgun / YaHei, etc.) so non-Latin text is preserved. Without a Unicode font,
text is latin-1-safe (unsupported codepoints become '?'); prefer DOCX for
full Unicode fidelity in that case.
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Literal

from engine.storage.transcript_segments_repository import TranscriptSegmentRow
from engine.stt.speaker_voice_profile import resolve_speaker_label

ExportBinaryFormat = Literal["pdf", "docx"]

logger = logging.getLogger(__name__)

# Overridable in tests (point at a known-good TTF).
_UNICODE_FONT_PATH_OVERRIDE: Path | None = None

_UNICODE_FONT_CANDIDATES: tuple[Path, ...] = (
    Path(r"C:/Windows/Fonts/arialuni.ttf"),
    Path(r"C:/Windows/Fonts/Arial Unicode MS.ttf"),
    Path(r"C:/Windows/Fonts/malgun.ttf"),
    Path(r"C:/Windows/Fonts/msyh.ttc"),
    Path(r"C:/Windows/Fonts/YuGothR.ttc"),
    Path(r"C:/Windows/Fonts/segoeui.ttf"),
    Path(r"C:/Windows/Fonts/arial.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
)


def _pdf_safe(text: str) -> str:
    """Latin-1 fallback when no Unicode font is available."""
    return text.encode("latin-1", "replace").decode("latin-1")


def _resolve_unicode_font_path() -> Path | None:
    if _UNICODE_FONT_PATH_OVERRIDE is not None:
        return _UNICODE_FONT_PATH_OVERRIDE if _UNICODE_FONT_PATH_OVERRIDE.is_file() else None
    for path in _UNICODE_FONT_CANDIDATES:
        if path.is_file():
            return path
    return None


def _configure_pdf_unicode_font(pdf: object) -> bool:
    """Register a Unicode TTF on ``pdf`` as family ``OmniUnicode``. True if ok."""
    font_path = _resolve_unicode_font_path()
    if font_path is None:
        return False
    try:
        pdf.add_font("OmniUnicode", "", str(font_path))  # type: ignore[attr-defined]
        pdf.add_font("OmniUnicode", "B", str(font_path))  # type: ignore[attr-defined]
        return True
    except Exception:
        logger.warning("could not load Unicode PDF font from %s", font_path, exc_info=True)
        return False


def _pdf_set_font(pdf: object, *, unicode_ok: bool, bold: bool = False, size: float = 11) -> None:
    style = "B" if bold else ""
    if unicode_ok:
        pdf.set_font("OmniUnicode", style, size)  # type: ignore[attr-defined]
    else:
        pdf.set_font("Helvetica", style, size)  # type: ignore[attr-defined]


def _pdf_cell(pdf: object, text: str, *, unicode_ok: bool, line_height: float) -> None:
    body = text if unicode_ok else _pdf_safe(text)
    pdf.multi_cell(pdf.epw, line_height, body)  # type: ignore[attr-defined]


def _pdf_write_lines(
    pdf: object, lines: list[str], *, unicode_ok: bool, line_height: float = 6
) -> None:
    for line in lines:
        if line:
            _pdf_cell(pdf, line, unicode_ok=unicode_ok, line_height=line_height)
        else:
            pdf.ln(3)  # type: ignore[attr-defined]


def _plain_lines_from_markdown(markdown: str) -> list[str]:
    lines: list[str] = []
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            lines.append("")
            continue
        if line.startswith("#"):
            lines.append(re.sub(r"^#+\s*", "", line))
            continue
        lines.append(line)
    return lines


def export_transcript_pdf(segments: list[TranscriptSegmentRow], identity: str = "Me") -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    unicode_ok = _configure_pdf_unicode_font(pdf)
    _pdf_set_font(pdf, unicode_ok=unicode_ok, size=11)
    for segment in segments:
        speaker = resolve_speaker_label(segment.speaker_id, identity)
        line = f"{speaker}: {segment.text}"
        _pdf_cell(pdf, line, unicode_ok=unicode_ok, line_height=6)
        pdf.ln(2)
    return bytes(pdf.output())


def export_transcript_docx(segments: list[TranscriptSegmentRow], identity: str = "Me") -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Meeting transcript", level=1)
    for segment in segments:
        speaker = resolve_speaker_label(segment.speaker_id, identity)
        document.add_paragraph(f"{speaker}: {segment.text}")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_meeting_pdf(
    title: str,
    notes_text: str,
    enhanced_notes_md: str,
    segments: list[TranscriptSegmentRow],
    identity: str = "Me",
) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    unicode_ok = _configure_pdf_unicode_font(pdf)
    _pdf_set_font(pdf, unicode_ok=unicode_ok, bold=True, size=16)
    _pdf_cell(pdf, title, unicode_ok=unicode_ok, line_height=8)
    pdf.ln(4)
    _pdf_set_font(pdf, unicode_ok=unicode_ok, size=11)
    if enhanced_notes_md.strip():
        _pdf_set_font(pdf, unicode_ok=unicode_ok, bold=True, size=13)
        _pdf_cell(pdf, "Enhanced Notes", unicode_ok=unicode_ok, line_height=7)
        _pdf_set_font(pdf, unicode_ok=unicode_ok, size=11)
        _pdf_write_lines(
            pdf, _plain_lines_from_markdown(enhanced_notes_md), unicode_ok=unicode_ok
        )
        pdf.ln(2)
    if notes_text.strip():
        _pdf_set_font(pdf, unicode_ok=unicode_ok, bold=True, size=13)
        _pdf_cell(pdf, "My Notes", unicode_ok=unicode_ok, line_height=7)
        _pdf_set_font(pdf, unicode_ok=unicode_ok, size=11)
        _pdf_write_lines(
            pdf,
            [line.strip() for line in notes_text.splitlines() if line.strip()],
            unicode_ok=unicode_ok,
        )
        pdf.ln(2)
    if segments:
        _pdf_set_font(pdf, unicode_ok=unicode_ok, bold=True, size=13)
        _pdf_cell(pdf, "Transcript", unicode_ok=unicode_ok, line_height=7)
        _pdf_set_font(pdf, unicode_ok=unicode_ok, size=11)
        for segment in segments:
            speaker = resolve_speaker_label(segment.speaker_id, identity)
            _pdf_cell(
                pdf,
                f"{speaker}: {segment.text}",
                unicode_ok=unicode_ok,
                line_height=6,
            )
            pdf.ln(2)
    return bytes(pdf.output())


def export_meeting_docx(
    title: str,
    notes_text: str,
    enhanced_notes_md: str,
    segments: list[TranscriptSegmentRow],
    identity: str = "Me",
) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(title, level=0)
    if enhanced_notes_md.strip():
        document.add_heading("Enhanced Notes", level=1)
        for line in _plain_lines_from_markdown(enhanced_notes_md):
            if line:
                document.add_paragraph(line)
    if notes_text.strip():
        document.add_heading("My Notes", level=1)
        for line in notes_text.splitlines():
            if line.strip():
                document.add_paragraph(line.strip())
    if segments:
        document.add_heading("Transcript", level=1)
        for segment in segments:
            speaker = resolve_speaker_label(segment.speaker_id, identity)
            document.add_paragraph(f"{speaker}: {segment.text}")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
